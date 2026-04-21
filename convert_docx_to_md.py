"""Convert all DOCX files in doc/ directory to Markdown format."""
import os
from pathlib import Path
from docx import Document


def convert_table_to_md(table) -> str:
    """Convert a docx table to markdown format."""
    if not table.rows:
        return ""

    lines = []

    # Header row
    header_cells = [cell.text.strip().replace('\n', ' ') for cell in table.rows[0].cells]
    lines.append("| " + " | ".join(header_cells) + " |")
    lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")

    # Data rows
    for row in table.rows[1:]:
        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        lines.append("| " + " | ".join(cells) + " |")

    return "\n".join(lines)


def docx_to_md(docx_path: Path, md_path: Path):
    """Convert a single DOCX file to Markdown."""
    print(f"Converting: {docx_path.name}")

    doc = Document(str(docx_path))

    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(f"# {docx_path.stem}\n\n")

        for element in doc.element.body:
            # Handle paragraphs
            if element.tag.endswith('p'):
                # Find corresponding paragraph object
                for para in doc.paragraphs:
                    if para._element is element:
                        text = para.text.strip()
                        if text:
                            # Check if it's a heading based on style
                            style_name = para.style.name if para.style else ""
                            if "Heading 1" in style_name:
                                f.write(f"# {text}\n\n")
                            elif "Heading 2" in style_name:
                                f.write(f"## {text}\n\n")
                            elif "Heading 3" in style_name:
                                f.write(f"### {text}\n\n")
                            elif "Heading 4" in style_name:
                                f.write(f"#### {text}\n\n")
                            else:
                                f.write(f"{text}\n\n")
                        break

            # Handle tables
            elif element.tag.endswith('tbl'):
                for table in doc.tables:
                    if table._element is element:
                        f.write(convert_table_to_md(table) + "\n\n")
                        break

    print(f"  -> Saved: {md_path.name}")


def main():
    doc_dir = Path(__file__).parent / "doc"
    output_dir = doc_dir / "markdown"
    output_dir.mkdir(exist_ok=True)

    docx_files = list(doc_dir.glob("*.docx"))

    if not docx_files:
        print("No DOCX files found in doc/ directory")
        return

    print(f"Found {len(docx_files)} DOCX files\n")

    for docx_path in docx_files:
        md_path = output_dir / f"{docx_path.stem}.md"
        try:
            docx_to_md(docx_path, md_path)
        except Exception as e:
            print(f"  Error converting {docx_path.name}: {e}")

    print(f"\nDone! Markdown files saved to: {output_dir}")


if __name__ == "__main__":
    main()
